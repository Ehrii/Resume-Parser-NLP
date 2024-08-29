from flask import Flask, request, redirect, url_for, render_template, flash, send_file,jsonify
from werkzeug.utils import secure_filename
import os
import docx
import MySQLdb
from config import Config
from datetime import datetime
import re
import spacy
from spacy.pipeline import EntityRuler
from spacy.language import Language
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
import io
from spacy.lang.en.stop_words import STOP_WORDS
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL

# Load spaCy model
nlp = spacy.load('en_core_web_sm')

# Create an EntityRuler
@Language.factory("education_entity_ruler")
def create_education_ruler(nlp, name):
    ruler = EntityRuler(nlp)
    patterns = [
        {"label": "EDUCATION", "pattern": "Pre-school"},
        {"label": "EDUCATION", "pattern": "Kindergarten"},
        {"label": "EDUCATION", "pattern": "Elementary"},
        {"label": "EDUCATION", "pattern": [{"LOWER": "primary"}, {"LOWER": "school"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "junior"}, {"LOWER": "high"}, {"LOWER": "school"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "secondary"}, {"LOWER": "school"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "senior"}, {"LOWER": "high"}, {"LOWER": "school"}]},
        {"label": "EDUCATION", "pattern": "TVET"},
        {"label": "EDUCATION", "pattern": "Technical-Vocational"},
        {"label": "EDUCATION", "pattern": "TESDA"},
        {"label": "EDUCATION", "pattern": [{"LOWER": "associate's"}, {"LOWER": "degree"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "associate"}, {"LOWER": "degree"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "some"}, {"LOWER": "college"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "college"}, {"LOWER": "graduate"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "vocational"}, {"LOWER": "graduate"}]},
        {"label": "EDUCATION", "pattern": [{"LOWER": "vocational"}, {"LOWER": "undergraduate"}]},
        {"label": "EDUCATION", "pattern": "Bachelor"},
        {"label": "EDUCATION", "pattern": [{"LOWER": "bachelor's"}, {"LOWER": "degree"}]},
        {"label": "EDUCATION", "pattern": "Master"},
        {"label": "EDUCATION", "pattern": [{"LOWER": "master's"}, {"LOWER": "degree"}]},
        {"label": "EDUCATION", "pattern": "Doctoral"},
        {"label": "EDUCATION", "pattern": [{"LOWER": "doctoral"}, {"LOWER": "degree"}]},
        {"label": "EDUCATION", "pattern": "PhD"},
        {"label": "EDUCATION", "pattern": [{"LOWER": "doctor"}, {"LOWER": "of"}, {"LOWER": "philosophy"}]},
    ]
    ruler.add_patterns(patterns)
    return ruler

# Add the EntityRuler to the pipeline
if 'education_entity_ruler' not in nlp.pipe_names:
    nlp.add_pipe('education_entity_ruler', last=True)
else:
    print("EntityRuler already in the pipeline.")

# Save the spaCy model to disk
nlp.to_disk("trained_spacy_model")

app = Flask(__name__)
app.config.from_object(Config)

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

db = MySQLdb.connect(
    host=app.config['MYSQL_HOST'],
    user=app.config['MYSQL_USER'],
    passwd=app.config['MYSQL_PASSWORD'],
    db=app.config['MYSQL_DB']
)

# Define the path to the skills.txt file
skills_file_path = os.path.join(os.path.dirname(__file__), 'linkedin_skills.txt')

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def clean_text(text):
    text = text.replace('\n', ' ')  # Replace newlines with spaces
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with a single space
    text = text.strip()  # Remove leading and trailing spaces
    return text

def preprocess_text(text):
    # Replace multiple spaces with a single space
    text = re.sub(r'\s+', ' ', text)
    
    # Add newlines before common resume section headers
    section_headers = ['OBJECTIVE', 'ACHIEVEMENTS', 'EDUCATIONAL BACKGROUND', 'SKILLS', 'PERSONAL DATA', 'CHARACTER REFERENCE']
    for header in section_headers:
        text = re.sub(f'({header})', r'\n\n\1\n', text)
    
    # Add newlines after colons
    text = re.sub(r':([^\n])', r':\n\1', text)
    
    # Separate items in lists
    text = re.sub(r'([a-z])([A-Z])', r'\1\n\2', text)
    
    # Add newlines before dates
    text = re.sub(r'([a-zA-Z])(\d{1,2}/\d{1,2}/\d{4})', r'\1\n\2', text)
    
    # Add newlines before phone numbers
    text = re.sub(r'(\d{11})', r'\n\1', text)
    
    # Add newlines before and after email addresses
    text = re.sub(r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Z|a-z]{2,})', r'\n\1\n', text)
    
    # Remove extra newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

def extract_text_from_pdf(pdf_path):
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
    page_interpreter = PDFPageInterpreter(resource_manager, converter)

    with open(pdf_path, 'rb') as fh:
        for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
            page_interpreter.process_page(page)
            text = fake_file_handle.getvalue()
            yield text

    converter.close()
    fake_file_handle.close()

@app.route('/', methods=['GET', 'POST'])
@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    resume_tip_message = {
        'title': 'Optimize Your Job Matching',
        'content': [
            "To ensure the most accurate job matching results, we recommend using our standardized resume format.",
            "This format is designed to highlight your skills and experiences in a way that our system can easily interpret, leading to better job recommendations.",
            "Key benefits of using our format:",
            [
                "Improved accuracy in skill and experience recognition",
                "Higher chances of matching with suitable job opportunities",
                "Consistent presentation of your qualifications"
            ],
            "Download our resume template and enhance your job matching potential today!"
        ],
        'cta': {
            'text': 'Download Resume Template',
            'link': '/download_template_docx'
        }
    }

    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return render_template('upload.html', resume_tip_message=resume_tip_message)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file')
            return render_template('upload.html', resume_tip_message=resume_tip_message)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_type = filename.rsplit('.', 1)[1].lower()
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            text_content = parse_resume(filepath)
            age = extract_age(text_content)
            skills, skills_warning = extract_skills(text_content)
            education = extract_education(text_content)
            upload_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            resume_id = filename, file_type, upload_date, text_content, age, skills, education

            data = {
            "text_content": text_content, 
            "filename": filename,
            "age":age, 
            "skills":skills, 
            "education":education,
            }
            return jsonify(data)

            # Find suitable jobs
            # matched_jobs = find_matching_jobs(resume_id, age, skills, education)
            
            # Clean up the file after processing
            # os.remove(filepath)
            
            # warning = None
            # no_match_message = None
            # if isinstance(matched_jobs, dict):
            #     if 'error' in matched_jobs:
            #         warning = matched_jobs['error']
            #         matched_jobs = []
            #     elif 'no_match' in matched_jobs:
            #         no_match_message = matched_jobs['no_match']
            #         matched_jobs = []
            
            # skills_warning_message = None
            # if skills_warning:
            #     skills_warning_message = "If you believe that you included skills in your resume, they may not be recognized by our current database. These skills may not be compatible with the jobs in our system, which could affect your job matches. Consider reviewing and updating your skills to align with industry-standard terminology."
            
            # return render_template('upload.html', 
            #                        text_content=text_content, 
            #                        formatted_text=text_content.replace('\n', '<br>'),
            #                        age=age, 
            #                        skills=skills, 
            #                        education=education, 
            #                        matched_jobs=matched_jobs,
            #                        warning=warning,
            #                        no_match_message=no_match_message,
            #                        skills_warning_message=skills_warning_message,
            #                        resume_tip_message=resume_tip_message)
        else:
            flash('Allowed file types are pdf, doc, docx')
            return render_template('upload.html', resume_tip_message=resume_tip_message)
    
    # This return statement is for when the method is GET
    return render_template('upload.html', resume_tip_message=resume_tip_message)

def parse_resume(filepath):
    if filepath.endswith('.pdf'):
        raw_text = ' '.join(extract_text_from_pdf(filepath))
    elif filepath.endswith('.doc') or filepath.endswith('.docx'):
        raw_text = parse_doc(filepath)
    else:
        return ''
    
    return preprocess_text(raw_text)

def parse_doc(filepath):
    document = docx.Document(filepath)
    text = ''
    for para in document.paragraphs:
        text += para.text + '\n'
    
    return text.strip()

def extract_age(text):
    birthdate_patterns = [
        r'\b(?:born|birthdate|dob|date of birth)[\s:]*([A-Za-z]+ \d{1,2}, \d{4})\b',  # Example: January 01, 2003
        r'\b(?:born|birthdate|dob|date of birth)[\s:]*([\d]{1,2}/[\d]{1,2}/[\d]{4})\b',  # Example: 01/01/2003
        r'\b(?:born|birthdate|dob|date of birth)[\s:]*([\d]{4}-[\d]{2}-[\d]{2})\b'  # Example: 2003-01-01
    ]

    age_patterns = [
        r'\bAge[\s:]*([0-9]+)\b',  
        r'\b([0-9]+)\s*years?\s*old\b',  
        r'\b([0-9]+)\+\b',  
    ]

    # Check for birthdates first
    for pattern in birthdate_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            birthdate_str = matches[0]
            try:
                for fmt in ('%B %d, %Y', '%m/%d/%Y', '%Y-%m-%d'):
                    try:
                        birthdate = datetime.strptime(birthdate_str, fmt)
                        today = datetime.today()
                        age = today.year - birthdate.year - ((today.month, today.day) < (birthdate.month, birthdate.day))
                        return age
                    except ValueError:
                        continue
            except ValueError:
                pass

    # If no birthdate found, check for direct age mentions
    for pattern in age_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))

    return None

def load_skills_from_file(filename):
    try:
        with open(filename, 'r') as file:
            skills = [line.strip() for line in file if line.strip()]
        return skills
    except FileNotFoundError:
        print(f"File {filename} not found.")
        return []

def update_skills_file_from_db():
    cursor = db.cursor()
    cursor.execute("SELECT job_skills FROM job_descriptions")
    job_skills_results = cursor.fetchall()
    cursor.close()

    # Extract skills from the database and convert to lowercase
    new_skills = set()
    for job_skills in job_skills_results:
        skills_list = [skill.strip().lower() for skill in job_skills[0].split(',')]
        new_skills.update(skills_list)

    # Load existing skills from the file
    existing_skills = load_skills_from_file(skills_file_path)
    existing_skills_set = set(skill.lower() for skill in existing_skills)

    # Combine both sets to avoid duplication
    updated_skills = existing_skills_set.union(new_skills)

    # Save updated skills back to the file
    try:
        with open(skills_file_path, 'w') as file:
            for skill in sorted(updated_skills):
                file.write(f"{skill}\n")
        print("Skills.txt updated successfully.")
    except Exception as e:
        print(f"Error updating skills.txt: {e}")

# Call the update function at an appropriate place in your code
update_skills_file_from_db()

def extract_skills(text):
    skill_list = load_skills_from_file(skills_file_path)

    skills_section = re.search(r'\b(?:SKILLS?|ABILITIES|COMPETENCIES)\b', text, re.IGNORECASE)
    
    if not skill_list:
        return 'No skills found', skills_section is not None
    
    text_lower = text.lower()
    skills_found = []
    
    for skill in skill_list:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text_lower):
            skills_found.append(skill)
    
    return ', '.join(skills_found) if skills_found else 'No skills found', skills_section is not None and not skills_found


def normalize_education(edu):
    edu = edu.lower()
    if any(term in edu for term in ['pre-school', 'preschool', 'kindergarten']):
        return 'Early Childhood Education'
    elif any(term in edu for term in ['elementary', 'primary']):
        return 'Elementary School'
    elif any(term in edu for term in ['junior high', 'secondary']):
        return 'Junior High School'
    elif 'senior high' in edu:
        return 'Senior High School'
    elif 'associate' in edu:
        return 'Associate\'s Degree'
    elif 'some college' in edu:
        return 'Some College Level'
    elif 'college graduate' in edu:
        return 'College Graduate'
    elif 'vocational graduate' in edu:
        return 'Vocational Graduate'
    elif 'vocational undergraduate' in edu:
        return 'Vocational Undergraduate'
    elif any(term in edu for term in ['bachelor', 'college', 'university']):
        return 'Bachelor\'s Degree'
    elif 'master' in edu:
        return 'Master\'s Degree'
    elif any(term in edu for term in ['phd', 'doctorate', 'doctoral']):
        return 'Doctoral Degree'
    elif any(term in edu for term in ['tvet', 'technical-vocational', 'tesda']):
        return 'Technical-Vocational Education and Training'
    else:
        return edu.title()

def extract_education(text):
    doc = nlp(text)
    education_info = [ent.text for ent in doc.ents if ent.label_ == 'EDUCATION']
    
    # Additional rule-based extraction
    education_keywords = [
        'pre-school', 'kindergarten', 'elementary', 'primary', 'primary school',
        'junior high', 'secondary', 'secondary school', 'senior high',
        'tvet', 'technical-vocational', 'tesda', 'associate degree', 'bachelor',
        'college', 'university', 'master', 'doctoral', 'phd', 'some college',
        'college graduate', 'vocational graduate', 'vocational undergraduate'
    ]
    
    # Use regex to find education keywords
    for keyword in education_keywords:
        matches = re.findall(r'\b' + re.escape(keyword) + r'\b', text.lower())
        education_info.extend(matches)
    
    # Normalize and remove duplicates
    normalized_education = list(set(normalize_education(edu) for edu in education_info))
    
    # Sort education levels
    education_order = [
        'Early Childhood Education', 'Elementary School', 'Junior High School',
        'Senior High School', 'Vocational Undergraduate', 'Technical-Vocational Education and Training',
        'Some College Level', 'Associate\'s Degree', 'Vocational Graduate',
        'Bachelor\'s Degree', 'College Graduate', 'Master\'s Degree', 'Doctoral Degree'
    ]
    sorted_education = sorted(normalized_education, key=lambda x: education_order.index(x) if x in education_order else len(education_order))
    
    return ', '.join(sorted_education) if sorted_education else 'No education information found'

# def store_resume(file_name, file_type, upload_date, text_content, age, skills, education):
#     cursor = db.cursor()
#     query = """
#         INSERT INTO resumes (file_name, file_type, upload_date, text_content, age, skills, education) 
#         VALUES (%s, %s, %s, %s, %s, %s, %s)
#     """
#     cursor.execute(query, (file_name, file_type, upload_date, text_content, age, skills, education))
#     resume_id = cursor.lastrowid
#     db.commit()
#     cursor.close()
#     print("Resume stored in database.")
    
#     return resume_id

# def find_matching_jobs(resume_id, age, skills, education):
#     missing_info = []

#     if education is None or education == 'No education information found':
#         missing_info.append("educational attainment")
#     if age is None:
#         missing_info.append("age")
#     if skills is None or skills == 'No skills found':
#         missing_info.append("skills")

#     if missing_info:
#         missing_info_str = ", ".join(missing_info)
#         return {
#             "error": f"Unable to match jobs. The following information is missing from your resume: {missing_info_str}. Please ensure your resume includes these details for accurate job matching."
#         }
    
#     cursor = db.cursor()
#     query = "SELECT job_title, job_requirements, job_skills, age_requirement FROM job_descriptions"
#     cursor.execute(query)
#     jobs = cursor.fetchall()
#     cursor.close()

#     matched_jobs = []
    
#     # Define education hierarchy
#     education_hierarchy = {
#         'Doctoral Degree': 11,
#         'Master\'s Degree': 10,
#         'College Graduate': 9,
#         'Bachelor\'s Degree': 8,
#         'Vocational Graduate': 7,
#         'Associate\'s Degree': 6,
#         'Some College Level': 5,
#         'Vocational Undergraduate': 4,
#         'Technical-Vocational Education and Training': 3,
#         'Senior High School': 2,
#         'Junior High School': 1,
#         'Elementary School': 0
#     }

#     def get_education_level(edu):
#         for key, value in education_hierarchy.items():
#             if key.lower() in edu.lower():
#                 return value, key
#         return -1, None 

#     for job in jobs:
#         job_title, job_requirements, job_skills, age_requirement = job

#         # Age matching
#         age_match = True
#         if age_requirement:
#             age_requirement = age_requirement.strip()
#             if '-' in age_requirement:
#                 try:
#                     min_age, max_age = map(int, age_requirement.split('-'))
#                     if not (min_age <= age <= max_age):
#                         age_match = False
#                 except ValueError:
#                     age_match = False
#             else:
#                 try:
#                     required_age = int(age_requirement.replace('+', '').strip())
#                     if age < required_age:
#                         age_match = False
#                 except ValueError:
#                     age_match = False

#         # Skill matching
#         skill_matches = [skill for skill in skills.split(', ') if skill.lower() in job_skills.lower()]
        
#         # Education matching
#         job_required_education_level = -1
#         required_education = None
#         for edu_level, value in education_hierarchy.items():
#             if edu_level.lower() in job_requirements.lower():
#                 job_required_education_level = value
#                 required_education = edu_level
#                 break

#         if job_required_education_level == -1:
#             continue

#         matching_education_levels = [
#             edu for edu in education.split(', ')
#             if get_education_level(edu)[0] >= job_required_education_level and get_education_level(edu)[0] != -1
#         ]

#         if matching_education_levels and skill_matches and age_match:
#             matched_jobs.append({
#                 'job_title': job_title,
#                 'job_requirements': job_requirements,
#                 'job_skills': job_skills,
#                 'age_requirement': age_requirement,
#                 'skill_matches': ', '.join(skill_matches),
#                 'required_education': required_education,
#                 'matching_education_levels': ', '.join(matching_education_levels),
#                 'applicant_highest_education': max(education.split(', '), key=lambda x: get_education_level(x)[0]),
#                 'skill_match_count': len(skill_matches)  # Add a count of matched skills
#             })

#             # Store each match
#             store_job_match(resume_id, job_title)

#     if not matched_jobs:
#         return {
#             "no_match": "We couldn't find any job matches at the moment based on your current qualifications. This doesn't mean you're not qualified, just that we don't have suitable openings right now. Consider expanding your skills or checking back later for new opportunities."
#         }
    
#     # Sort matched jobs by the number of skill matches, from most to least
#     matched_jobs.sort(key=lambda x: x['skill_match_count'], reverse=True)

#     return matched_jobs

    
# def store_job_match(resume_id, job_title):
#     cursor = db.cursor()
#     query = """
#         INSERT INTO job_matches (resume_id, job_title) 
#         VALUES (%s, %s)
#     """
#     cursor.execute(query, (resume_id, job_title))
#     db.commit()
#     cursor.close()
#     print("Job match stored in database.")

@app.route('/download_template_docx')
def download_template_docx():
    document = Document()

    # Add content to the document
    document.add_heading('FULL NAME', 0)
    
    # Create a table for layout (2 columns: text and image)
    table = document.add_table(rows=1, cols=2)
    table.allow_autofit = False
    table.columns[0].width = Inches(5) 
    table.columns[1].width = Inches(1.5)  

    # Add contact information to the left cell
    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = left_cell.paragraphs[0]
    p.add_run('Email: your.email@example.com\n').bold = True
    p.add_run('Phone: (123) 456-7890\n').bold = True
    p.add_run('Address: City, State').bold = True

    # Add placeholder image to the right cell
    right_cell = table.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = right_cell.paragraphs[0]
    placeholder_image_path = os.path.join(app.root_path, 'static', 'placeholder_image.png')
    if os.path.exists(placeholder_image_path):
        run = p.add_run()
        run.add_picture(placeholder_image_path, width=Inches(1.5), height=Inches(1.5))

    # Objective
    document.add_heading('OBJECTIVE', level=1)
    document.add_paragraph('A brief statement about your career goals and what you are looking for.')

    # Education
    document.add_heading('EDUCATION', level=1)
    p = document.add_paragraph()
    p.add_run('Degree, Major\n').bold = True
    p.add_run('University Name, Graduation Year')

    # Skills
    document.add_heading('SKILLS', level=1)
    p = document.add_paragraph()
    p.add_run('Hard Skills: ').bold = True
    p.add_run('Skill 1, Skill 2, Skill 3\n')
    p.add_run('Soft Skills: ').bold = True
    p.add_run('Skill 1, Skill 2, Skill 3')

    # Work Experience
    document.add_heading('WORK EXPERIENCE', level=1)
    p = document.add_paragraph()
    p.add_run('Job Title, Company Name\n').bold = True
    p.add_run('Start Date - End Date\n').italic = True
    document.add_paragraph('Responsibility/Achievement 1', style='List Bullet')
    document.add_paragraph('Responsibility/Achievement 2', style='List Bullet')

    # Certifications
    document.add_heading('CERTIFICATIONS', level=1)
    document.add_paragraph('Certification Name, Issuing Organization, Year', style='List Bullet')

    # Additional Information
    document.add_heading('ADDITIONAL INFORMATION', level=1)
    document.add_paragraph('Languages, Volunteer Work, etc.')

    # Save the document to a BytesIO object
    f = io.BytesIO()
    document.save(f)
    f.seek(0)

    return send_file(f, as_attachment=True, download_name='resume_template.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == "__main__":
    app.run(debug=True)

